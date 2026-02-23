import docx

def analyze_docx(path):
    try:
        doc = docx.Document(path)
        print(f"Analyzing {path}...")
        
        # Check tables
        print(f"\nFound {len(doc.tables)} tables.")
        
        for i, table in enumerate(doc.tables):
            print(f"\n--- Table {i+1} ---")
            # Print first few rows to understand structure
            for r_idx, row in enumerate(table.rows[:3]):
                cells = [cell.text.strip() for cell in row.cells]
                print(f"Row {r_idx}: {cells}")
                
        # Check paragraphs for context
        print(f"\n--- First 5 Paragraphs ---")
        for p in doc.paragraphs[:5]:
            if p.text.strip():
                print(p.text.strip())
                
    except Exception as e:
        print(f"Error reading docx: {e}")

if __name__ == "__main__":
    analyze_docx("resultApar.docx")
