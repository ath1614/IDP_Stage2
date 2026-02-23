import argparse
import requests
import json
import docx
import time
import os
# import weasyprint  # Only needed for Disciplinary PDF generation
import markdown
from dotenv import load_dotenv
from run_ocr import run_ocr, pdf_to_base64_images

# Load environment variables
project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
load_dotenv(os.path.join(project_root, ".env"))

# Configuration
OCR_VM_IP = os.getenv("OCR_VM_IP", "34.47.176.38")
LLM_VM_IP = os.getenv("LLM_VM_IP", "34.93.108.135")

OCR_VM_URL = f"http://{OCR_VM_IP}:8000"
LLM_VM_URL = f"http://{LLM_VM_IP}:8000"

REQUIRED_FIELDS = """
You are an expert AI Assistant specialized in extracting data from Indian Government Annual Performance Appraisal Reports (APAR).
Your task is to analyze the provided OCR text and extract specific performance details into a structured JSON format.

### EXTRACTION RULES:

1. **Officer Details**:
   - **officer_name**: Full name of the officer being appraised. Check header/footer or first page. Consistency is key.
   - **date_of_birth**: Date of birth (DD/MM/YYYY).

2. **APAR Entries (Year-wise)**:
   - Identify each assessment year (e.g., "2018-19", "2019-2020", "01.04.2019 to 31.03.2020").
   - For EACH year found, extract the following details. If a year is split across pages, merge the info.
   
     - **Reporting Authority**:
       - `name`: Name of the Reporting Officer.
       - `grading`: Numerical score (1-10). If "Outstanding" infer 10, "Very Good" infer 8, "Good" infer 6.
       - `pen_picture`: Extract the **FULL** descriptive assessment (Pen Picture) by the Reporting Officer. 
         - **CRITICAL**: Do not summarize. Capture the exact text about "general assessment", "quality of work", "state of health", etc.
         - If text is > 500 chars, you may condense slightly but keep key adjectives.

     - **Reviewing Authority**:
       - `name`: Name of the Reviewing Officer.
       - `grading`: Numerical score (1-10).
         - *Note*: If the Reviewing Officer says "I agree with Reporting Officer" without a separate grade, infer the same grade.
         - *Important*: If the text says "Reviewing Officer did not supervise for minimum period" (e.g. < 3 months), set `grading` to null.
       - `pen_picture`: Comments by the Reviewing Officer. Capture full remarks on agreement/disagreement.

     - **Accepting Authority**:
       - `name`: Name of the Accepting Authority.
       - `grading`: Numerical score (1-10) given by the Accepting Authority.
       - `remarks`: Final remarks or decision (e.g., "Fit for promotion", "Grading accepted").

     - **Integrity**:
       - Look for the specific "Integrity" column or certification.
       - **Default**: "Beyond Doubt".
       - **Exception**: Only change this if the text explicitly mentions "Doubtful", "Complaint pending", or "Under observation".

### JSON OUTPUT STRUCTURE:
Return ONLY this valid JSON object:
{
  "officer_name": "Name String",
  "date_of_birth": "DD/MM/YYYY",
  "apar_entries": [
    {
      "year": "YYYY-YYYY",
      "reporting": {
        "name": "Name String",
        "grading": "Value or null",
        "pen_picture": "Text content"
      },
      "reviewing": {
        "name": "Name String",
        "grading": "Value or null",
        "pen_picture": "Text content"
      },
      "accepting": {
        "name": "Name String",
        "grading": "Value or null",
        "remarks": "Text content"
      },
      "integrity": "Beyond Doubt"
    }
  ]
}
"""

DISCIPLINARY_SUMMARY_PROMPT = """
You are an expert Legal AI Assistant specialized in summarizing Indian Government Disciplinary Proceedings.
Your task is to analyze the provided OCR text from a Disciplinary Case file and generate a comprehensive Executive Summary in structured JSON format.

### INSTRUCTIONS:
1. **Analyze the Content**: Read the document carefully to understand the context, charges, proceedings, and timeline.
2. **Structure the Report**: Generate a professional report following the CNN/DailyMail benchmark style.
3. **Systematic Flow**: Ensure the report presents a logical, chronological flow of events.
4. **Robustness & Partial Matches**:
   - If specific fields are missing or the document is incomplete, **DO NOT FAIL**.
   - Extract **WHATEVER YOU CAN** from the available text.
   - If a field cannot be found, explicitly state "Not explicitly mentioned in the provided text" or infer from context if possible.
   - **NEVER** return an empty result. Even a partial summary is better than nothing.
5. **Detail**: Include as much specific detail (dates, names, rule numbers, quote snippets) as possible in the summary fields.

### REQUIRED JSON STRUCTURE:
Return a valid JSON object with the following keys:
{
  "headline": "A formal title summarizing the case",
  "executive_summary": "A detailed high-level overview (approx. 200-300 words) of the entire case. Include all available context.",
  "charged_officer_details": {
    "name": "Name of the officer (or 'Unknown')",
    "designation": "Designation (or 'Unknown')",
    "details": "Other relevant details"
  },
  "background_and_chronology": "A detailed systematic account of events leading to the proceedings. Include dates and letter references.",
  "key_allegations": ["List of specific charges or allegations. Be comprehensive."],
  "evidence_and_findings": "Key evidence presented and findings of the inquiry officer. Include specific document references if available.",
  "defense_arguments": "Summary of the defense provided by the charged officer. If not present, state 'Defense arguments not included in this document'.",
  "conclusion_status": "The final outcome or current status of the case."
}

### OUTPUT RULES:
- The content within the JSON fields should be well-written, professional, and detailed (approx 1.5 - 2 pages total content).
- **CRITICAL**: Return ONLY the JSON object. Do not add markdown formatting outside the JSON.
- **CRITICAL**: Ensure the JSON is valid. Escape quotes within strings.
"""

CLASSIFICATION_PROMPT = """
You are a Document Classifier.
Analyze the following text and classify it into exactly one of these categories:
1. "APAR" (Annual Performance Appraisal Report) - Look for keywords: "Appraisal", "Reporting Officer", "Reviewing Officer", "Integrity".
2. "DISCIPLINARY" (Disciplinary Case) - Look for keywords: "Article of Charge", "Memorandum", "Inquiry", "Imputation of Misconduct", "Show Cause".
3. "OTHER" - If it doesn't fit the above.

Return ONLY a JSON object:
{
  "doc_type": "APAR" | "DISCIPLINARY" | "OTHER",
  "confidence": 0.0 to 1.0
}
"""

import os

# ... imports ...

def classify_document(text_sample):
    """Classify the document type based on the first few pages of text."""
    print("Classifying document type...")
    
    # Keyword-based Pre-check (Fast & Reliable)
    text_lower = text_sample.lower()
    
    # Disciplinary Keywords
    disciplinary_keywords = [
        "article of charge", "memorandum", "disciplinary proceeding", 
        "inquiry report", "imputation of misconduct", "show cause notice",
        "prosecution brief", "charged officer", "inquiry authority"
    ]
    
    # APAR Keywords
    apar_keywords = [
        "annual performance appraisal", "reporting officer", "reviewing officer",
        "accepting authority", "pen picture", "integrity certificate"
    ]
    
    disc_score = sum(1 for k in disciplinary_keywords if k in text_lower)
    apar_score = sum(1 for k in apar_keywords if k in text_lower)
    
    print(f"Keyword Scores - Disciplinary: {disc_score}, APAR: {apar_score}")
    
    # Strong signal override
    if disc_score > 2 and disc_score > apar_score:
        print(">> Classified as DISCIPLINARY based on keywords.")
        return "DISCIPLINARY"
    if apar_score > 2 and apar_score > disc_score:
        print(">> Classified as APAR based on keywords.")
        return "APAR"

    # Fallback to LLM if ambiguous
    print("Keywords ambiguous, asking LLM...")
    prompt = CLASSIFICATION_PROMPT + f"\n\nDocument Sample:\n{text_sample[:3000]}"
    
    payload = {
        "text": text_sample[:3000],
        "schema": CLASSIFICATION_PROMPT 
    }
    
    try:
        response = requests.post(f"{LLM_VM_URL}/api/extract", json=payload, timeout=120)
        if response.status_code == 200:
            res = response.json().get("result", "{}")
            if "```json" in res:
                res = res.split("```json")[1].split("```")[0]
            elif "```" in res:
                res = res.split("```")[1].split("```")[0]
            data = json.loads(res.strip())
            return data.get("doc_type", "OTHER")
    except Exception as e:
        print(f"Classification failed: {e}")
        
    print(">> Defaulting to APAR")
    return "APAR" # Default fallback

def summarize_chunk(chunk_text, chunk_index, total_chunks):
    """Summarize a single chunk of a large document."""
    print(f"  Summarizing chunk {chunk_index+1}/{total_chunks}...")
    prompt = f"""
    You are analyzing Part {chunk_index+1} of {total_chunks} of a Disciplinary Case document.
    Summarize the key information in this section, focusing on:
    - New allegations or charges mentioned.
    - Key evidence or witness statements.
    - Important procedural steps or dates.
    
    Keep the summary concise but informative (approx. 300 words).
    If this section contains no significant information (e.g., just headers or noise), reply with "No significant information."
    
    CRITICAL: Do NOT return JSON. Return plain text only.
    """
    return _call_llm(prompt + "\n\n" + chunk_text, schema=None, expect_json=False)

def recursive_summarize(full_text, depth=0):
    # Base case: If text is small enough, summarize directly
    # Lowered threshold to 12000 chars (~3000 tokens) to avoid context limit issues
    if len(full_text) <= 12000:
        combined_text = DISCIPLINARY_SUMMARY_PROMPT + "\n\n" + full_text
        return _call_llm(combined_text, schema=None, expect_json=True)
    
    print(f"    [Depth {depth}] Text too long ({len(full_text)} chars). Splitting...")
    
    # Split into chunks
    CHUNK_SIZE = 10000  # Reduced chunk size for safety
    overlap = 500
    step = CHUNK_SIZE - overlap
    chunks = [full_text[i:i+CHUNK_SIZE] for i in range(0, len(full_text), step)]
    chunk_summaries = []
    
    # Process chunks
    # Note: recursive_summarize call inside summarize_chunk needs to handle depth if we were fully recursive,
    # but here we just do one level of chunking then combine.
    # Actually, the original logic was flat chunking. 
    # Let's keep it simple: Map (summarize chunks) -> Reduce (summarize summaries)
    
    # We need to make sure we don't infinitely recurse if chunks don't shrink enough, 
    # but summarization always shrinks text significantly.
    
    for i, chunk in enumerate(chunks):
        print(f"    [Depth {depth}] Processing chunk {i+1}/{len(chunks)}...")
        summary = summarize_chunk(chunk, i, len(chunks))
        chunk_summaries.append(summary)
    
    combined_summary = "\n\n".join(chunk_summaries)
    print(f"    [Depth {depth}] Combined summary length: {len(combined_summary)} chars. Recursively summarizing...")
    
    # Recursive call in case the combined summary is still too large
    return recursive_summarize(combined_summary, depth=depth+1)

def summarize_disciplinary_case(full_text):
    """Summarize disciplinary case using LLM"""
    print("Generating Disciplinary Summary...")
    return recursive_summarize(full_text)

def extract_fields_with_llm(page_texts, mode="auto"):
    """
    Send text to LLM VM and get JSON back
    mode: "auto" (classify), "apar" (force APAR), "summary" (force Disciplinary Summary)
    """
    
    # Combine all pages into one text block with explicit page markers
    # This ensures the LLM understands the physical flow of the document
    full_text_parts = []
    for idx, text in enumerate(page_texts):
        if text.strip():
            full_text_parts.append(f"--- Page {idx+1} ---\n{text}")
            
    full_text = "\n\n".join(full_text_parts)
    
    # 1. Determine Document Type
    doc_type = "OTHER"
    
    if mode == "auto":
        doc_type = classify_document(full_text)
        print(f"Detected Document Type: {doc_type}")
    elif mode == "summary":
        print("Mode is 'summary'. Forcing DISCIPLINARY processing.")
        doc_type = "DISCIPLINARY"
    elif mode == "apar":
        print("Mode is 'apar'. Forcing APAR processing.")
        doc_type = "APAR"
    else:
        print(f"Unknown mode '{mode}'. Defaulting to APAR.")
        doc_type = "APAR"
    
    if doc_type == "DISCIPLINARY" or doc_type == "OTHER":
        print(f"Treating {doc_type} as DISCIPLINARY for summarization.")
        
        # Check if text is empty or too short
        if not full_text or len(full_text.strip()) < 100:
            print("Warning: Text is too short for meaningful summarization.")
            return {
                "type": "DISCIPLINARY",
                "data": "Error: Document text is empty or too short."
            }
            
        return {
            "type": "DISCIPLINARY",
            "data": summarize_disciplinary_case(full_text)
        }
    
    # Else, assume APAR (existing logic)
    
    # Chunking Strategy:
    # If text > 18000 chars, we split it.
    MAX_CHARS = 18000 
    
    final_data = None
    
    if len(full_text) <= MAX_CHARS:
        print(f"Sending full text ({len(full_text)} chars) to LLM at {LLM_VM_URL}...")
        res = _call_llm(full_text)
        
        if isinstance(res, str):
            res = _extract_json_from_text(res) or {}
            
        if not isinstance(res, dict):
             res = {}
             
        res["type"] = "APAR" # Tag it
        final_data = res
    else:
        print(f"Text too long ({len(full_text)} chars). Splitting into chunks...")
        chunks = [full_text[i:i+MAX_CHARS] for i in range(0, len(full_text), MAX_CHARS)]
        
        merged_data = {
            "type": "APAR",
            "officer_name": "", 
            "date_of_birth": "", 
            "apar_entries": []
        }
        
        for idx, chunk in enumerate(chunks):
            print(f"  Processing chunk {idx+1}/{len(chunks)}...")
            # Add context to help LLM understand this is a partial document
            chunk_prompt = f"Note: This is part {idx+1} of a larger document.\n\n{chunk}"
            result = _call_llm(chunk_prompt)
            
            if isinstance(result, str):
                result = _extract_json_from_text(result) or {}
                
            if not isinstance(result, dict):
                print(f"    Warning: Chunk {idx+1} result is not valid JSON. Skipping.")
                continue
            
            # Merge logic
            if not merged_data["officer_name"] and result.get("officer_name"):
                merged_data["officer_name"] = result["officer_name"]
            if not merged_data["date_of_birth"] and result.get("date_of_birth"):
                merged_data["date_of_birth"] = result["date_of_birth"]
                
            if result.get("apar_entries"):
                merged_data["apar_entries"].extend(result["apar_entries"])
                
        final_data = merged_data

    # Save JSON for benchmarking
    # We don't have output_dir context here easily unless we pass it or assume it.
    # The caller usually handles output, but for now let's return it.
    return final_data

def _call_llm(text, schema=REQUIRED_FIELDS, expect_json=True):
    payload = {
        "text": text
    }
    if schema:
        payload["schema"] = schema
    
    for attempt in range(3):
        try:
            # If retry attempt (attempt > 0), try simplifying to text-only if JSON is failing
            current_expect_json = expect_json
            if attempt > 1 and expect_json:
                print("    > Attempting fallback to plain text extraction (dropping schema constraint)...")
                if "schema" in payload:
                    del payload["schema"]
                current_expect_json = False

            response = requests.post(f"{LLM_VM_URL}/api/extract", json=payload, timeout=300)
            if response.status_code == 200:
                try:
                    resp_json = response.json()
                    result = resp_json.get("result", "{}")
                    
                    # Cleanup potential markdown code blocks
                    if "```json" in result:
                        result = result.split("```json")[1].split("```")[0]
                    elif "```" in result:
                        result = result.split("```")[1].split("```")[0]
                    
                    result = result.strip()

                    # Check for empty or single-brace artifacts
                    # Allow short results if we are not expecting JSON (e.g. "No info")
                    if expect_json:
                        if not result or result == "{" or len(result) < 5:
                             print(f"    Attempt {attempt+1}: Result truncated or empty ('{result}'). Retrying...")
                             time.sleep(2)
                             continue
                    else:
                        # For plain text, allow anything non-empty. 
                        # But if it looks like "{}" (empty json) and we didn't want json, treat as empty?
                        if not result or result == "{}":
                             # It might be that the LLM defaulted to returning empty JSON even when we asked for text
                             # In that case, we can probably treat it as "No information found" instead of retrying forever
                             if result == "{}":
                                 print(f"    Attempt {attempt+1}: Received empty JSON '{{}}' in text mode. Treating as empty text.")
                                 return ""
                             print(f"    Attempt {attempt+1}: Result empty. Retrying...")
                             time.sleep(2)
                             continue

                    if not current_expect_json:
                        # If we fell back to text, return text. 
                        # But if the caller *really* wanted JSON (expect_json=True), we might try to wrap it?
                        # For now, let's just return the text and let _extract_json_from_text handle it later if possible
                        return result
                        
                    return json.loads(result)
                except json.JSONDecodeError:
                    if not current_expect_json:
                         # Plain text mode, return raw result
                         return result
                         
                    print(f"    Attempt {attempt+1}: JSON Decode Error. Raw response: {response.text[:200]}...")
                    time.sleep(2)
                    continue
            else:
                print(f"    Attempt {attempt+1}: LLM Error {response.status_code} - {response.text[:200]}")
                time.sleep(2)
                continue
        except Exception as e:
            print(f"    Attempt {attempt+1}: Request Failed: {e}")
            time.sleep(2)
            
    print("LLM Extraction failed after 3 attempts.")
    return {} if expect_json else ""

def _extract_json_from_text(text):
    """
    Robustly extract the first valid JSON object from a string by counting braces.
    Handles cases with prefix text, suffix text, or multiple JSONs (takes first).
    """
    if not isinstance(text, str):
        return None
        
    start_index = text.find('{')
    if start_index == -1:
        return None
        
    stack = 0
    for i in range(start_index, len(text)):
        if text[i] == '{':
            stack += 1
        elif text[i] == '}':
            stack -= 1
            if stack == 0:
                json_candidate = text[start_index:i+1]
                try:
                    return json.loads(json_candidate)
                except json.JSONDecodeError:
                    # Continue searching if this one failed (unlikely if brace count matches, but possible with quotes)
                    # Actually if brace count matches but invalid json, it's probably issue with content
                    # But let's be safe and return None or try harder? 
                    # For now, return None is safer than crashing
                    return None
    return None

def generate_disciplinary_docx(data, output_dir, filename_base):
    """Generate the Disciplinary Executive Summary Report"""
    # Ensure filename ends with .docx
    filename = f"{filename_base}_Summary.docx"
    output_path = os.path.join(output_dir, filename)
    
    try:
        doc = docx.Document()
        
        # Style Definitions
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = docx.shared.Pt(11)
        
        # Title
        title = doc.add_paragraph("EXECUTIVE SUMMARY: DISCIPLINARY CASE")
        title.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].bold = True
        title.runs[0].font.size = docx.shared.Pt(18)
        title.runs[0].font.underline = True
        title.runs[0].font.color.rgb = docx.shared.RGBColor(0, 51, 102) # Dark Blue
        
        doc.add_paragraph() # Spacer
        
        lines = []
        json_data = None
        
        if isinstance(data, dict):
            json_data = data
        elif isinstance(data, str):
            # Check for empty JSON artifact
            if data.strip() == "{}" or data.strip() == '"{}"' or data.strip() == "'{}'":
                doc.add_paragraph("Summary generation returned empty result.")
                doc.save(output_path)
                return output_path
            
            # Attempt robust JSON extraction
            json_data = _extract_json_from_text(data)
                
        if json_data:
            # If valid JSON, convert to formatted text manually
            # We expect keys like officer_name, case_reference, allegations, etc.
            # Or maybe just "executive_summary"
            
            # Construct a Markdown-like string from the dict
            md_lines = []
            
            # Headline
            headline = json_data.get('headline') or json_data.get('case_reference') or json_data.get('Headline') or 'Disciplinary Case'
            md_lines.append(f"# {headline}")
            
            # Officer Details
            officer_info = json_data.get('charged_officer_details') or json_data.get('officer_name') or json_data.get('Charged Officer Details') or 'N/A'
            if isinstance(officer_info, dict):
                officer_str = ", ".join([f"{k.replace('_', ' ').title()}: {v}" for k, v in officer_info.items()])
            else:
                officer_str = str(officer_info)
            md_lines.append(f"**Charged Officer Details:** {officer_str}")
            md_lines.append("")
            
            # Background & Chronology
            background = json_data.get('background_and_chronology') or json_data.get('Background & Chronology') or json_data.get('Background')
            if background:
                    md_lines.append("## Background & Chronology")
                    md_lines.append(str(background))
                    md_lines.append("")

            # Executive Summary
            exec_sum = json_data.get('executive_summary') or json_data.get('Executive Summary')
            if exec_sum:
                    md_lines.append("## Executive Summary")
                    md_lines.append(exec_sum)
                    md_lines.append("")
                    
            # Allegations
            allegations = json_data.get('key_allegations') or json_data.get('allegations') or json_data.get('Key Allegations')
            if allegations:
                md_lines.append("## Key Allegations")
                if isinstance(allegations, list):
                    for alg in allegations:
                        md_lines.append(f"- {alg}")
                else:
                    md_lines.append(str(allegations))
                md_lines.append("")

            # Evidence & Findings
            evidence = json_data.get('evidence_and_findings') or json_data.get('inquiry_findings') or json_data.get('Evidence & Findings')
            if evidence:
                md_lines.append("## Evidence & Findings")
                if isinstance(evidence, dict):
                    for k, v in evidence.items():
                        md_lines.append(f"**{k.replace('_', ' ').title()}:** {v}")
                elif isinstance(evidence, list):
                        for item in evidence:
                            md_lines.append(f"- {item}")
                else:
                    md_lines.append(str(evidence))
                md_lines.append("")
            
            # Defense Arguments
            defense = json_data.get('defense_arguments') or json_data.get('Defense Arguments')
            if defense:
                md_lines.append("## Defense Arguments")
                md_lines.append(str(defense))
                md_lines.append("")
                
            # Conclusion
            conclusion = json_data.get('conclusion_status') or json_data.get('conclusion') or json_data.get('Conclusion/Status')
            if conclusion:
                md_lines.append("## Conclusion/Status")
                md_lines.append(str(conclusion))
            
            # Use the constructed lines
            lines = md_lines
        else:
            # Markdown Parsing Mode
            lines = str(data).split('\n')

                
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            if line.startswith('# '):
                # H1
                h = doc.add_heading(line[2:], level=1)
                h.runs[0].font.color.rgb = docx.shared.RGBColor(0, 0, 0)
            elif line.startswith('## '):
                # H2
                h = doc.add_heading(line[3:], level=2)
                h.runs[0].font.color.rgb = docx.shared.RGBColor(50, 50, 50)
            elif line.startswith('### '):
                # H3
                h = doc.add_heading(line[4:], level=3)
            elif line.startswith('- ') or line.startswith('* '):
                # Bullet
                p = doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith('**') and line.endswith('**'):
                # Bold Line
                p = doc.add_paragraph()
                run = p.add_run(line[2:-2])
                run.bold = True
            else:
                # Normal Text
                # Handle bolding within text **bold**
                p = doc.add_paragraph()
                parts = line.split('**')
                for i, part in enumerate(parts):
                    run = p.add_run(part)
                    if i % 2 == 1: # Odd parts are between ** **
                        run.bold = True
        
        doc.save(output_path)
        print(f"Disciplinary Report generated at {output_path}")
        return output_path
        
    except Exception as e:
        print(f"Error generating disciplinary docx: {e}")
        import traceback
        traceback.print_exc()
        return None


def generate_apar_docx(data, output_dir, filename_base="APAR_single_table"):
    """
    Generate the APAR docx table per strict complex structure.
    Rows: Header Row 1, Header Row 2, Data Rows (5 years).
    Cols: 9 Columns total.
    Col 1: Name
    Col 2: DoB
    Col 3: Year
    Col 4: Rep Grading
    Col 5: Rep Pen Picture
    Col 6: Rev Grading
    Col 7: Rev Pen Picture
    Col 8: Acc Grading
    Col 9: Acc Pen Picture
    """
    # Ensure filename ends with .docx
    if not filename_base.lower().endswith(".docx"):
        filename = f"{filename_base}.docx"
    else:
        filename = filename_base
        
    output_path = os.path.join(output_dir, filename)
    
    def safe_str(val):
        """Safely convert value to string, handling None/Null"""
        if val is None:
            return ""
        s = str(val).strip()
        if s.lower() in ["null", "none", ""]:
            return ""
        return s

    try:
        doc = docx.Document()
        
        # Set Landscape
        section = doc.sections[0]
        section.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height
        section.left_margin = docx.shared.Inches(0.5)
        section.right_margin = docx.shared.Inches(0.5)
        
        # Create Table: 7 Rows (2 Header + 5 Data), 9 Cols
        table = doc.add_table(rows=7, cols=9)
        table.style = 'Table Grid'
        
        # --- HEADER ROW 1 (Index 0) ---
        row0 = table.rows[0].cells
        
        # Name (Col 0) - Merge with Row 1
        row0[0].text = "Name of Officer"
        row0[0].merge(table.rows[1].cells[0])
        
        # DoB (Col 1) - Merge with Row 1
        row0[1].text = "Date of Birth (DoB)"
        row0[1].merge(table.rows[1].cells[1])
        
        # Year (Col 2) - Merge with Row 1
        row0[2].text = "APAR Year"
        row0[2].merge(table.rows[1].cells[2])
        
        # Reporting (Cols 3-4)
        row0[3].text = "Reporting Officer"
        row0[3].merge(row0[4])
        
        # Reviewing (Cols 5-6)
        row0[5].text = "Reviewing Officer"
        row0[5].merge(row0[6])
        
        # Accepting (Cols 7-8)
        row0[7].text = "Accepting Authority"
        row0[7].merge(row0[8])
        
        # Style Row 0
        for cell in row0:
            p = cell.paragraphs[0]
            p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            if p.runs:
                p.runs[0].font.bold = True
                p.runs[0].font.size = docx.shared.Pt(11)

        # --- HEADER ROW 2 (Index 1) ---
        row1 = table.rows[1].cells
        # Cols 0, 1, 2 are already merged from above
        
        # Sub-headers for Authorities
        sub_headers = ["Grading", "Pen Picture Analysis"] * 3 # Rep, Rev, Acc
        
        # Indices in Row 1 corresponding to Cols 3-8
        for i, text in enumerate(sub_headers):
            col_idx = 3 + i
            row1[col_idx].text = text
            p = row1[col_idx].paragraphs[0]
            p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            if p.runs:
                p.runs[0].font.bold = True
                p.runs[0].font.size = docx.shared.Pt(10)
                
        # --- DATA ROWS (Indices 2-6) ---
        # We have 5 slots for years
        entries = data.get("apar_entries", [])
        
        # Ensure we have exactly 5 entries (pad or truncate)
        # Sort by year if possible? Assuming chronological order from LLM
        display_entries = entries[:5]
        while len(display_entries) < 5:
            display_entries.append({})
            
        # 1. Fill Officer Details (Spanning 5 rows)
        # Name (Col 0): Merge Rows 2-6
        c_name_start = table.rows[2].cells[0]
        c_name_end = table.rows[6].cells[0]
        c_name_start.merge(c_name_end)
        c_name_start.text = safe_str(data.get("officer_name"))
        c_name_start.vertical_alignment = docx.enum.table.WD_ALIGN_VERTICAL.CENTER
        
        # DoB (Col 1): Merge Rows 2-6
        c_dob_start = table.rows[2].cells[1]
        c_dob_end = table.rows[6].cells[1]
        c_dob_start.merge(c_dob_end)
        c_dob_start.text = safe_str(data.get("date_of_birth"))
        c_dob_start.vertical_alignment = docx.enum.table.WD_ALIGN_VERTICAL.CENTER
        
        # 2. Fill Year and Data (Row by Row)
        for i, entry in enumerate(display_entries):
            row_idx = 2 + i
            row_cells = table.rows[row_idx].cells
            
            # Year (Col 2)
            # User asked for "1, 2, 3, 4, 5" numbers, OR actual year. 
            # "Display APAR Year numbers: 1 2 3 4 5"
            # But later "For each APAR Year (1 to 5)... Provide values"
            # I will put the Actual Year String, but if empty, maybe just the index number?
            # Let's put the Actual Year. If specific requirement for "1", "2" comes back, we change.
            # Actually user said: "Display APAR Year numbers: 1 2 3 4 5" literally.
            # But that might mean the count.
            # I'll put the Year String if present, else just i+1.
            year_text = safe_str(entry.get("year"))
            if not year_text and i < 5:
                 # If no data, maybe just leave empty? Or put number?
                 # Let's just use the number as fallback or prefix
                 pass 
            
            # Let's trust the "APAR Year" column header usually implies the actual year text.
            row_cells[2].text = year_text if year_text else f"{i+1}"
            row_cells[2].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            
            # Reporting
            rep = entry.get("reporting", {}) or {}
            row_cells[3].text = safe_str(rep.get("grading"))
            row_cells[3].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            row_cells[4].text = safe_str(rep.get("pen_picture"))
            
            # Reviewing
            rev = entry.get("reviewing", {}) or {}
            rev_grad = safe_str(rev.get("grading"))
            if not rev_grad:
                 # Check for the "did not supervise" case
                 rev_pen = safe_str(rev.get("pen_picture"))
                 if "did not supervise" in rev_pen.lower():
                     row_cells[3].text = "-"
            row_cells[5].text = rev_grad
            row_cells[5].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            row_cells[6].text = safe_str(rev.get("pen_picture"))

            # Accepting
            acc = entry.get("accepting", {}) or {}
            row_cells[7].text = safe_str(acc.get("grading"))
            row_cells[7].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            row_cells[8].text = safe_str(acc.get("remarks")) # "Pen Picture Analysis" col for Accepting matches "Remarks"
            
            # Center vertically
            for c_idx in range(2, 9):
                # Vertical alignment isn't easily exposed in python-docx for all cell types without OXML
                # But we can try setting the property if possible, or just leave as top
                pass

        doc.save(output_path)
        print(f"Document generated at {output_path}")
        return output_path
        
    except Exception as e:
        print(f"Error generating docx: {e}")
        import traceback
        traceback.print_exc()
        return None

def generate_disciplinary_pdf(data, output_dir, filename_base):
    """Generate PDF report using WeasyPrint"""
    filename = f"{filename_base}_Summary.pdf"
    output_path = os.path.join(output_dir, filename)
    
    # 1. Prepare Content (HTML)
    json_data = None
    if isinstance(data, dict):
        json_data = data
    elif isinstance(data, str):
        if data.strip() == "{}" or data.strip() == '"{}"':
             return None 
        json_data = _extract_json_from_text(data)

    if not json_data:
        print(f"DEBUG: generate_disciplinary_pdf received empty json_data. Input type: {type(data)}. Value: {str(data)[:100]}...")
        # Fallback for plain text
        html_content = f"<h1>Disciplinary Summary</h1><p>{str(data)}</p>"
    else:
        # Construct HTML from JSON
        headline = json_data.get('headline') or 'Disciplinary Case'
        
        # Officer
        officer_info = json_data.get('charged_officer_details') or 'N/A'
        officer_html = ""
        if isinstance(officer_info, dict):
            officer_html = "<ul>" + "".join([f"<li><b>{k.replace('_', ' ').title()}:</b> {v}</li>" for k,v in officer_info.items()]) + "</ul>"
        else:
            officer_html = f"<p>{officer_info}</p>"
            
        # Sections
        sections_html = ""
        # Handle flexible keys
        key_map = {
            'background_and_chronology': 'Background & Chronology',
            'executive_summary': 'Executive Summary',
            'key_allegations': 'Key Allegations',
            'evidence_and_findings': 'Evidence & Findings',
            'defense_arguments': 'Defense Arguments',
            'conclusion_status': 'Conclusion/Status'
        }
        
        # Also check for capitalized keys
        
        for key, title in key_map.items():
            val = json_data.get(key) or json_data.get(title) or json_data.get(key.replace('_', ' ').title())
            if val:
                sections_html += f"<h2>{title}</h2>"
                if isinstance(val, list):
                    sections_html += "<ul>" + "".join([f"<li>{item}</li>" for item in val]) + "</ul>"
                elif isinstance(val, dict):
                     sections_html += "<ul>" + "".join([f"<li><b>{k.replace('_', ' ').title()}:</b> {v}</li>" for k,v in val.items()]) + "</ul>"
                else:
                    # Check for markdown content (simple bold/bullets)
                    # Convert markdown to HTML if needed, but for now just text
                    val_str = str(val).replace('\n', '<br>')
                    sections_html += f"<p>{val_str}</p>"
        
        html_content = f"""
        <html>
        <head>
            <style>
                body {{ font-family: 'Calibri', sans-serif; font-size: 11pt; line-height: 1.4; color: #333; }}
                h1 {{ color: #003366; text-decoration: underline; text-align: center; font-size: 18pt; margin-bottom: 20px; }}
                h2 {{ color: #333333; border-bottom: 1px solid #ccc; padding-bottom: 5px; margin-top: 20px; font-size: 14pt; }}
                ul {{ margin-left: 20px; }}
                li {{ margin-bottom: 5px; }}
                .officer-box {{ background-color: #f5f5f5; padding: 10px; border: 1px solid #ddd; border-radius: 5px; margin-bottom: 20px; }}
            </style>
        </head>
        <body>
            <h1>{headline}</h1>
            <div class="officer-box">
                <h3>Charged Officer Details</h3>
                {officer_html}
            </div>
            {sections_html}
        </body>
        </html>
        """
        
    try:
        # Requires weasyprint and system dependencies
        import weasyprint
        weasyprint.HTML(string=html_content).write_pdf(output_path)
        print(f"PDF generated at {output_path}")
        return output_path
    except ImportError:
        print("WeasyPrint not available. Skipping PDF generation. Install with: brew install pango")
        return None
    except Exception as e:
        print(f"Error generating PDF: {e}")
        return None

def main(pdf_path, custom_output_dir=None, mode="auto"):
    # 1. Run OCR
    print(f"Step 1: Running OCR on {pdf_path}...")
    try:
        images = pdf_to_base64_images(pdf_path)
        page_texts = [""] * len(images) # Initialize list to store text per page
        batch_size = 3
        
        print(f"Total pages: {len(images)}. Processing in batches of {batch_size}...")
        
        for i in range(0, len(images), batch_size):
            batch = images[i:i+batch_size]
            print(f"  Sending batch {i//batch_size + 1} (pages {i+1}-{i+len(batch)})...")
            
            payload = {"images": batch, "languages": ["en"]}
            
            # Add retry logic
            for attempt in range(3):
                try:
                    ocr_resp = requests.post(f"{OCR_VM_URL}/api/ocr/batch", json=payload, timeout=300)
                    if ocr_resp.status_code == 200:
                        ocr_data = ocr_resp.json()
                        # Store text for each page in the correct index
                        for j, res in enumerate(ocr_data["results"]):
                            page_texts[i+j] = res["text"]
                        break
                    else:
                        print(f"    Batch failed (Status {ocr_resp.status_code}). Retrying...")
                        time.sleep(2)
                except Exception as e:
                     print(f"    Batch error: {e}. Retrying...")
                     time.sleep(2)
            else:
                print(f"    Failed to process batch {i//batch_size + 1} after 3 attempts.")

        # Verify all pages processed
        missing_pages = [i+1 for i, t in enumerate(page_texts) if not t]
        if missing_pages:
            print(f"Warning: No text extracted for pages: {missing_pages}")
        
        full_text_len = sum(len(p) for p in page_texts)
        print(f"OCR Complete. Extracted {full_text_len} characters from {len(page_texts)} pages.")
    except Exception as e:
        print(f"OCR Failed: {e}")
        return

    # 2. Run LLM Extraction
    print("Step 2: Extracting fields with LLM...")
    extracted_data = extract_fields_with_llm(page_texts, mode=mode)
    print("Extracted Data:", json.dumps(extracted_data, indent=2))
    
    # 3. Generate Document
    print("Step 3: Generating Word Document...")
    
    # Create output directory structure
    base_name = os.path.splitext(os.path.basename(pdf_path))[0]
    
    if custom_output_dir:
        final_output_dir = custom_output_dir
    else:
        # Structure: ./output/<PDF_NAME>/
        parent_output_dir = os.path.join(os.getcwd(), "output")
        final_output_dir = os.path.join(parent_output_dir, base_name)
    
    os.makedirs(final_output_dir, exist_ok=True)
    print(f"Output directory created: {final_output_dir}")

    # Check Document Type and Branch
    doc_type = extracted_data.get("type", "APAR")
    
    if doc_type == "DISCIPLINARY":
        print(">> Generating Disciplinary Executive Summary...")
        
        # Check data validity before passing
        disc_data = extracted_data.get("data", {})
        if not disc_data:
            print("WARNING: Disciplinary data is empty!")
            
        # Generate DOCX (always works)
        generate_disciplinary_docx(disc_data, final_output_dir, filename_base=base_name)
        
        # Try PDF if weasyprint available
        try:
            generate_disciplinary_pdf(disc_data, final_output_dir, filename_base=base_name)
        except:
            print("PDF generation skipped (weasyprint not available).")
    else:
        print(">> Generating APAR Data Sheet...")
        # Pass base_name as the desired filename
        generate_apar_docx(extracted_data, final_output_dir, filename_base=base_name)

if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("pdf_path")
    parser.add_argument("--output-dir", help="Custom output directory")
    parser.add_argument("--mode", default="auto", choices=["auto", "apar", "summary"], help="Processing mode")
    args = parser.parse_args()
    
    main(args.pdf_path, args.output_dir, args.mode)
